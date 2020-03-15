import pandas as pd
from docx import Document
from geopy import distance
import googlemaps
import re

gmaps = googlemaps.Client(key='AIzaSyB4XeimIWa4ArmXVyBE53HKv4GukiWQh9w')

def store_document_info(prop_df, agency_df, min_dist_to_service=2.0):
    agency_dict = agency_df.set_index('OBJECTID').to_dict()
    pd.set_option('mode.chained_assignment', None)
    new_col_dict = {'agency_address':'STREET_ADDRESS','agency_name':'AGENCY',
                    'agency_phone':'PHONE_NUMBER', 'agency_url':'WEBSITE_URL'}
    prop_df = prop_df[['opa_number','street_address', 'owner', 'total_due',
                        'lat','lon','closest agency']]
    prop_df['prop_coords'] = list(zip(prop_df.lat,prop_df.lon))

    for k, v in new_col_dict.items():
        prop_df[k] = prop_df.apply(lambda x: 
                    agency_dict[v][x['closest agency']], axis=1)

    prop_df['agency_coords'] = prop_df.apply(lambda x:
                            (agency_dict['Y'][x['closest agency']], 
                            agency_dict['X'][x['closest agency']]),
                            axis=1)
    prop_df['service_dist'] = prop_df.apply(lambda x: 
                            distance.distance(x.prop_coords,
                                x.agency_coords).miles,
                            axis=1)

    return prop_df[prop_df.service_dist > min_dist_to_service]


def get_directions(prop_coords, agency_coords, mode='driving'):
    orig = re.sub('[( )]','',str(prop_coords))
    end = re.sub('[( )]','',str(agency_coords))
    dir_obj = gmaps.directions(orig, end,mode=mode)
    total_duration = dir_obj[0]['legs'][0]['duration']['text']
    fare = 'Unknown'
    steps_list = []
    for i in dir_obj[0]['legs'][0]['steps']:
        string = i['html_instructions'] + ' ' + i['distance']['text']
        if 'transit_details' in i.keys():
            string += ' ({} for {} stops)'.format(
            i['transit_details']['line']['name'],
            i['transit_details']['num_stops'])
        cleaned_str = ' '.join(re.sub('<[^>]+>', ' ',
                    string).replace('&nbsp;', ' ').split())
        steps_list.append(cleaned_str)
    if mode == 'transit':
        if 'fare' in dir_obj[0].keys():
            fare = dir_obj[0]['fare']['text']
        return steps_list, total_duration, fare
    return steps_list, total_duration


def write_documents(df, path):
    for row in df.itertuples():
        driving, drive_dur = get_directions(row.prop_coords, row.agency_coords)
        transit, tran_dur, fare = get_directions(row.prop_coords,
            row.agency_coords, 'transit')
        dirs = get_directions(row.prop_coords, row.agency_coords)
        doc = Document()
        first_paragraph = ("You are receiving this notice because your "
            "property at {} is tax delinquent and we have identified that "
            "you have no housing counseling office in your immediate vicinity. "
            "Your nearest service location is {}. This is approximately {} "
            "mile(s) from your home and directions are enclosed at the bottom "
            "of this notice. While we offer services to help you remain in your "
            "home through this process and to prevent foreclosure, we partner "
            "with housing counseling agencies to offer important supplemental services.\
            ".format(row.street_address, row.agency_name, round(row.service_dist,2)))
    
        doc.add_paragraph(first_paragraph)
        doc.add_heading('Agency Information:', level=1)
        doc.add_paragraph(row.agency_name)
        doc.add_paragraph(row.agency_url)
        doc.add_paragraph(row.agency_phone)
        doc.add_paragraph(row.agency_address)

        doc.add_heading('Directions:', level=1)
        doc.add_paragraph('Driving from {} for about {}'.format(
            row.street_address, drive_dur),
            style='Intense Quote')
        if len(dirs) <= 1:
            doc.add_paragraph('Directions Unavailable')
        else:
            for step in driving:
                doc.add_paragraph(step, style='List Bullet')
        doc.add_paragraph('Transit from {} for about {}. FARES: {}'.format(
            row.street_address, tran_dur, fare),
            style='Intense Quote')
        if len(dirs) <= 1:
            doc.add_paragraph('Directions Unavailable')
        else:
            for step in transit:
                doc.add_paragraph(step, style='List Bullet')
        doc.save(path+'/'+str(row.opa_number)+'.docx')
